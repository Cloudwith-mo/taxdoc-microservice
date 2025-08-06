"""
Document Structure Extraction Service
Handles OCR, layout analysis, and text extraction for various document types
"""

import boto3
import json
import re
from typing import Dict, Any, List, Optional, Tuple
from docx import Document
import PyPDF2
from io import BytesIO

class DocumentStructureExtractor:
    def __init__(self):
        self.textract_client = boto3.client('textract', region_name='us-east-1')
    
    def extract_structure(self, file_bytes: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract document structure based on file type
        Returns structured text with layout information
        """
        if file_info['processing_route'] == 'ocr_pipeline':
            return self._extract_with_ocr(file_bytes, file_info)
        elif file_info['processing_route'] == 'text_extraction_pipeline':
            return self._extract_text_directly(file_bytes, file_info)
        else:
            raise ValueError(f"Unsupported processing route: {file_info['processing_route']}")
    
    def _extract_with_ocr(self, file_bytes: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Extract structure using OCR for images and scanned PDFs"""
        try:
            # Use Textract for OCR and layout analysis
            response = self.textract_client.analyze_document(
                Document={'Bytes': file_bytes},
                FeatureTypes=['TABLES', 'FORMS', 'LAYOUT']
            )
            
            # Parse Textract response into structured format
            structure = {
                'text_blocks': [],
                'tables': [],
                'key_value_pairs': [],
                'raw_text': '',
                'confidence_scores': [],
                'layout_info': {
                    'pages': 1,
                    'has_tables': False,
                    'has_forms': False
                }
            }
            
            for block in response.get('Blocks', []):
                if block['BlockType'] == 'LINE':
                    structure['text_blocks'].append({
                        'text': block['Text'],
                        'confidence': block.get('Confidence', 0) / 100.0,
                        'bbox': block.get('Geometry', {}).get('BoundingBox', {}),
                        'type': 'line'
                    })
                    structure['raw_text'] += block['Text'] + '\n'
                    structure['confidence_scores'].append(block.get('Confidence', 0) / 100.0)
                
                elif block['BlockType'] == 'TABLE':
                    structure['layout_info']['has_tables'] = True
                    # Extract table structure (simplified)
                    structure['tables'].append({
                        'id': block['Id'],
                        'confidence': block.get('Confidence', 0) / 100.0,
                        'bbox': block.get('Geometry', {}).get('BoundingBox', {})
                    })
                
                elif block['BlockType'] == 'KEY_VALUE_SET':
                    structure['layout_info']['has_forms'] = True
                    if block.get('EntityTypes') and 'KEY' in block['EntityTypes']:
                        structure['key_value_pairs'].append({
                            'key': block.get('Text', ''),
                            'confidence': block.get('Confidence', 0) / 100.0
                        })
            
            # Calculate overall confidence
            if structure['confidence_scores']:
                structure['overall_confidence'] = sum(structure['confidence_scores']) / len(structure['confidence_scores'])
            else:
                structure['overall_confidence'] = 0.0
            
            return structure
            
        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return self._create_error_structure(str(e))
    
    def _extract_text_directly(self, file_bytes: bytes, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text directly from text-based documents"""
        try:
            if file_info['category'] == 'docx':
                return self._extract_from_docx(file_bytes)
            elif file_info['category'] == 'pdf':
                return self._extract_from_pdf(file_bytes)
            elif file_info['category'] == 'text':
                return self._extract_from_text(file_bytes)
            else:
                raise ValueError(f"Unsupported text extraction for category: {file_info['category']}")
                
        except Exception as e:
            print(f"Text extraction failed: {e}")
            return self._create_error_structure(str(e))
    
    def _extract_from_docx(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        doc = Document(BytesIO(file_bytes))
        
        structure = {
            'text_blocks': [],
            'tables': [],
            'key_value_pairs': [],
            'raw_text': '',
            'confidence_scores': [1.0],  # High confidence for direct text
            'overall_confidence': 1.0,
            'layout_info': {
                'pages': 1,
                'has_tables': len(doc.tables) > 0,
                'has_forms': False
            }
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                structure['text_blocks'].append({
                    'text': para.text,
                    'confidence': 1.0,
                    'type': 'paragraph'
                })
                structure['raw_text'] += para.text + '\n'
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                table_text.append(row_text)
            
            structure['tables'].append({
                'data': table_text,
                'confidence': 1.0
            })
        
        return structure
    
    def _extract_from_pdf(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract text from PDF files"""
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        
        structure = {
            'text_blocks': [],
            'tables': [],
            'key_value_pairs': [],
            'raw_text': '',
            'confidence_scores': [0.9],  # Good confidence for PDF text
            'overall_confidence': 0.9,
            'layout_info': {
                'pages': len(pdf_reader.pages),
                'has_tables': False,
                'has_forms': False
            }
        }
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    structure['text_blocks'].append({
                        'text': text,
                        'confidence': 0.9,
                        'type': 'page',
                        'page_number': page_num + 1
                    })
                    structure['raw_text'] += text + '\n'
            except Exception as e:
                print(f"Failed to extract text from page {page_num}: {e}")
        
        return structure
    
    def _extract_from_text(self, file_bytes: bytes) -> Dict[str, Any]:
        """Extract from plain text files"""
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1', errors='ignore')
        
        return {
            'text_blocks': [{
                'text': text,
                'confidence': 1.0,
                'type': 'text'
            }],
            'tables': [],
            'key_value_pairs': [],
            'raw_text': text,
            'confidence_scores': [1.0],
            'overall_confidence': 1.0,
            'layout_info': {
                'pages': 1,
                'has_tables': False,
                'has_forms': False
            }
        }
    
    def _create_error_structure(self, error_msg: str) -> Dict[str, Any]:
        """Create error structure when extraction fails"""
        return {
            'text_blocks': [],
            'tables': [],
            'key_value_pairs': [],
            'raw_text': '',
            'confidence_scores': [],
            'overall_confidence': 0.0,
            'layout_info': {
                'pages': 0,
                'has_tables': False,
                'has_forms': False
            },
            'error': error_msg
        }
    
    def get_relevant_text_regions(self, structure: Dict[str, Any], max_tokens: int = 4000) -> str:
        """
        Extract most relevant text regions for LLM processing
        Prioritizes high-confidence, structured content
        """
        if 'error' in structure:
            return ""
        
        # Prioritize text blocks by confidence and type
        prioritized_blocks = sorted(
            structure['text_blocks'],
            key=lambda x: (x.get('confidence', 0), len(x['text'])),
            reverse=True
        )
        
        relevant_text = ""
        token_count = 0
        
        for block in prioritized_blocks:
            block_text = block['text']
            # Rough token estimation (1 token ≈ 4 characters)
            block_tokens = len(block_text) // 4
            
            if token_count + block_tokens > max_tokens:
                # Add partial text if it fits
                remaining_chars = (max_tokens - token_count) * 4
                if remaining_chars > 100:  # Only add if meaningful
                    relevant_text += block_text[:remaining_chars] + "..."
                break
            
            relevant_text += block_text + "\n"
            token_count += block_tokens
        
        return relevant_text.strip()